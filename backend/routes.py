from flask import Blueprint, request, jsonify,send_file,Response,current_app
from models import User,feedback,db,bcrypt,Question,Score,tz,Status
from flask_jwt_extended import create_access_token
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy import extract, func
from flask_cors import CORS,cross_origin
from flask_mail import Message
from datetime import datetime,timedelta
import random
import json,re
from extensions import mail 
from config import Config
from pypdf import PdfReader
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
import google.generativeai as genai
from reportlab.platypus import  SimpleDocTemplate,Paragraph,PageBreak,Spacer,Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from io import BytesIO
import traceback
from docx import Document
from calendar import month_abbr,monthrange



routes = Blueprint("routes", __name__)
CORS(routes)

otp_storage = {}


@routes.route("/Signup", methods=["POST"])
def Signup():
    data = request.get_json()
    username, email, password = data["username"], data["email"], data["password"]

    if User.query.filter_by(email=email).first():
        return jsonify({"error": "Email already exists"}), 400
    if User.query.filter_by(username=username).first():
        return jsonify({"error": "Username already exists"}), 400

    new_user = User(username=username, email=email)
    new_user.set_password(password)
    db.session.add(new_user)
    db.session.commit()

    return jsonify({"success": True, "message": "User registered successfully"}), 201


@routes.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    user = User.query.filter_by(email=data["email"]).first()

    if user and user.check_password(data["password"]):
        token = create_access_token(identity={"user_id":user.id})
        return jsonify({"success": True, "message": "Login successful", "token": token})

    return jsonify({"success": False, "error": "Invalid credentials"}), 401

# sending mail for otp
@routes.route('/send-otp', methods=['POST'])
def send_otp():
    data = request.json
    email = data.get("emailForReset")
    
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({"message": "Email not registered!"}), 400

    otp = str(random.randint(100000, 999999))
    otp_storage[email] = otp
    msg = Message("Password Reset OTP", sender=current_app.config["MAIL_USERNAME"], recipients=[email])
    msg.body = f"Your OTP for password reset is: {otp}"
    
    try:
        mail.send(msg)  
    except Exception as e:
        return jsonify({"error": "Failed to send email", "details": str(e)}), 500

    return jsonify({"message": "OTP sent successfully!"})

@routes.route('/verify-otp', methods=['POST'])
def verify_otp():
    data = request.json
    email = data.get("emailForReset")
    otp = data.get("otp")
    print(otp)
    print(otp_storage)
    if email not in otp_storage or otp_storage[email] != otp:
        return jsonify({"message": "Invalid OTP"}), 400
   
    del otp_storage[email]  
    return jsonify({"message": "OTP verified successfully"}), 200

# reset and update the password
@routes.route('/reset-password', methods=['POST'])
def reset_password():
    data=request.json 
    # reset-password
    if data["actions"]=="reset_password":
        email = data["emailForReset"]
        new_password=data["newPassword"]
        user = User.query.filter_by(email=email).first()
        """if not user:
            return jsonify({"message": "User not found!"}), 400"""
        hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
        user.password_hash= hashed_password
        db.session.commit()
    # update-password
    elif  data["actions"]=="update_password":
        id=data["user_id"]
        details=data["data"] 
        new_password = details["newPassword"]
        user=User.query.filter_by(id=id).first()
        old_password=details["currentPassword"]
        if not bcrypt.check_password_hash(user.password_hash,old_password):
            return jsonify({"message":"password does not match"}),400
    
        hashed=bcrypt.generate_password_hash(new_password).decode('utf-8')
        user.password_hash=hashed
        db.session.commit()
        
    else:
        return jsonify({"message":"something went wrong!!!"}),400


    return jsonify({"message": "Password reset successful!"}),200

    
@routes.route("/Pdffile", methods=["POST"])

def pdf_file():
    q_type={}
    if "pdfFile" not in request.files:
        return Response(f"No file uploaded",status=400,mimetype="text/plain")
    

    userfile = request.files["pdfFile"]
    user_filename = request.form.get("filename").lower
    difficulty=request.form.get("difficulty")
    questionCount=request.form.get("questionCount")

   

    try:
        all_text = ""
        if user_filename.endswith(".pdf"):
            readerpdf = PdfReader(userfile)
            if len(readerpdf.pages)==0:
                return Response(f"the file is empty",status=400,mimetype="text/plain")
            else:
                for i in range(len(readerpdf.pages)): 
                    page_text = readerpdf.pages[i].extract_text()
                    if page_text:
                        all_text += f"Page {i+1}:\n{page_text}\n\n"

        elif user_filename.endswith(".docx"):
            readerdocx=Document(userfile)
            if len(readerdocx.paragraphs)==0:
                return Response(f"the file is empty",status=400,mimetype="text/plain")
            else:
                for i in range(len(readerdocx.paragraphs)):
                    page_text=readerdocx.paragraphs[i].text


            


        # generating question using ai
        api_key = Config.API_KEY
        if not api_key:
            return Response(f"Error : API key is missing",status=401,mimetype="text/plain")
        try:
            
            genai.configure(api_key=api_key)
            model=genai.GenerativeModel("gemini-2.0-flash")
            model_response = model.generate_content(contents = (
                f"""
                Generate questions strictly based on the syllabus content provided below with {difficulty} to the syllabus
                Each question should be categorized by its corresponding mark allocation. 
                f"Ensure all questions are relevant to the syllabus and do not include any additional information. 
                Syllabus Content:{all_text}
                ### Response Format:
                Your response **must be** a valid JSON dictionary.
                Do **not** include any explanations, extra text, or formatting outside of JSON.
                Strictly follow those keys only:
                include the title also
                {{
                    "Questions":[{{"title":"Ai robotics"}}
                   {{ 
                    "marks":"2 mark",
                    "questions":[
                        " 1.what is robotics ?",
                         "2.what are recent innovation?"
                    ]
                    }},
                    {{
                    "marks":"2 mark",
                    "questions":[
                        " 1.what is ai ?",
                         "2.what is machine learning?"
                    ]
                    }}
                    ] 
                }}
                """
                ))

            response_text = model_response.candidates[0].content.parts[0].text
           # print(response_text)
            #print(type(response_text))
            clean_response = re.sub(r"```json\n|\n```", "", response_text).strip()
           # print(type(clean_response))
            try:
                model_output = json.loads(clean_response)
                print("Parsed Dictionary:", model_output)
            except json.JSONDecodeError as e:
                print("JSON Decode Error:", e)
                #print("Raw Response:", clean_response)

            
            #print(model_output)
            
            

            if not model_output :
                return Response("AI model did not generate any content", status=500, mimetype="text/plain")
        except Exception as e:
                return Response(f"Google AI Error: {str(e)}", status=500, mimetype="text/plain")

        # converting the text to document
        try:   
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()

            elements = []

            # Title
            title = model_output["Questions"][0].get("title", "No Title")
            elements.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
            elements.append(Spacer(1, 10))  
            left_style = ParagraphStyle(name="LeftAlign", parent=styles["Heading2"], alignment=TA_LEFT)
            right_style = ParagraphStyle(name="RightAlign", parent=styles["Heading2"], alignment=TA_RIGHT)

            for item, i in zip(model_output["Questions"][1:], newItems):
                data = [[
                Paragraph(f"<b>{i['marks']} marks</b>", left_style),  
                Paragraph(f"<b>{i['marks']} x {i['questions']} = {i['totalMarks']}</b>", right_style)  
            ]]

            # Create a table with two columns
                table = Table(data, colWidths=[200, 200])  # Adjust widths as needed
                table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, 0), 'LEFT'),    
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),  
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),  # Reduce spacing
                ]))


                elements.append(table)
                elements.append(Spacer(1, 10))
                if "questions" in item:
                    for question in item["questions"]:
                        elements.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{question}", styles["Normal"]))
                        elements.append(Spacer(1, 13))

                elements.append(Spacer(1, 12))  

            # Build PDF
            doc.build(elements)

            buffer.seek(0)
            return send_file(buffer, mimetype="application/pdf")

        except Exception as e:
            print(f"error:{e}")

    except Exception as e:
        print(f"Error: {e}") 
        return Response(f"Error {str(e)}",status=500 ,mimetype="text/plain") 
    
# feedback
@routes.route("/feedback",methods=["POST"])
def feedBack():
    data=request.json
    user_id=data['user_id']
    feedbacktype=data['feedbackType']
    feedback_message=data['feedbackText']
    current_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    user=User.query.filter_by(id=user_id).first()
    email_id=user.email
    username=user.username

    body=(f"NEW FEEDBACK RECEIVED \n"
          f"From:{email_id}\n"
          f"Time:{current_time}\n\n"
          f"feedback Type:{feedbacktype}\n\n"
          f"feedback from user:\n{feedback_message}"
          )
    # sending email
    msg=Message(subject=f"feedback from the {email_id}",
                sender=current_app.config["MAIL_USERNAME"],
                recipients=[current_app.config["MAIL_USERNAME"]],
                body=body)
    msg.reply_to=email_id
    try:
        mail.send(msg)
    except Exception as e:
        return jsonify({"message":f"{e} try back after some time"}),400
    
    # inserting the values in feedback table
    new_feedback=feedback(username=username,email=email_id,feedback_message=feedback_message)
    db.session.add(new_feedback)
    db.session.commit()
    
    #reply to the user
    return jsonify({"message": "Feedback sent successfully!"}), 200

@routes.route("/generate_assessment",methods=["POST"])
def generate_assessment():
    data=request.json
    user_id=data.get("user_id")
    collection=data.get("values")
    subject=collection.get("subject")
    topic=collection.get("topic")
    difficulty=collection.get("difficulty")
    num_of_quest=collection.get("questionCount")
    api_key=Config.API_KEY1
    if not api_key:
        return Response(f"Error : API key is missing",status=401,mimetype="text/plain")
    
    try:
        genai.configure(api_key=api_key)
        model=genai.GenerativeModel("gemini-2.0-flash")
        model_res=model.generate_content(contents=(
            f"""
            generate multiple choice question based on {subject}
            the topic related to the {topic} 
            generate question based on the level {difficulty}
            generate {num_of_quest} questions in given {subject}
            ### Response Format:
            Your response **must be** a valid JSON dictionary.
            Do **not** include any explanations, extra text, or formatting outside of JSON.
            Strictly follow those keys only:
            {{{{
                "question_text": "What is the capital of France?",
                "topic": "Geography",
                "choices": [
                {{"id":"a","text": "Paris"}},
                {{"id":"b","text": "Berlin"}},
                {{"id":"c","text": "London"}},
                {{"id":"d","text": "Madrid"}}
                ],
                "is_correct":"a" 
            }},
            {{
                "question_text": "Who invented Python?",
                "topic": "Programming",
                "choices": [
                {{"id":"a",text": "Guido van Rossum"}},
                {{"id":"b",text": "James Gosling"}},
                {{"id":"c","text": "Bjarne Stroustrup"}},
                {{"id":"d","text": "Dennis Ritchie"}}
                ],
                "is_correct": "a"
            }}}}
            """
        ))
        response_text = model_res.candidates[0].content.parts[0].text
        
        clean_response = re.sub(r"```json\n|\n```", "", response_text).strip()
        
        model_output = None
        try:
            model_output = json.loads(clean_response)
        except json.JSONDecodeError as e:
            print("JSON Decode Error:", e)

        
        
        if not model_output or not isinstance(model_output,list) :
                return Response("AI model did not generate any valid list of content", status=500, mimetype="text/plain")
        try:
            score=Score(subject=subject,topic=topic,difficulty=difficulty,user_id=user_id)
            db.session.add(score)
            db.session.flush() 
            # Get the score ID after flushing
            score_id = score.id
            for i in model_output:
                question=Question(score_id=score_id,quest_text=i["question_text"],choices=i["choices"],is_correct=i["is_correct"])
                db.session.add(question)
            db.session.commit()
        except SQLAlchemyError as e:
            db.session.rollback()
            return Response(f"An error occurred: {str(e)}", status=500, mimetype="text/plain")   
        
        return jsonify({
            "message": "Questions addded successfully!" 
        }), 200
    except Exception as e:
                return Response(f"Google AI Error: {str(e)}", status=500, mimetype="text/plain")
    
 # Starting the assessment   
@routes.route("/start",methods=["POST"])
def start():
    data=request.json
    user_id=data.get("user_id")
    id=data.get("score_id")

    #storing result
    score=Score.query.filter_by(user_id=user_id,id=id).first()
    if score is None:
        return jsonify({"error": "Assessment not found"}), 404
    result={"id":id,"subject":score.subject,"title":score.topic,"questions":[]}
    question=Question.query.filter_by(score_id=id)
    for q in question:
        options = []
        for c in q.choices:
            options.append({"id":c["id"],"option":c["text"]})
        question_data = {
            "id":q.id,
            "text": q.quest_text,
            "options": options
        }
        result["questions"].append(question_data)
    score.status="completed"
    db.session.commit()
    print(result)
    return jsonify(result)
    
#for pending requests
@routes.route("/pending",methods=["POST"])
def pending():
    data=request.json
    user_id=data["user_id"]
    #status=data["status"]
    user=User.query.filter_by(id=user_id).first()
    res=Score.query.filter_by(user_id=user_id).all()
    for r in res:
        if r.due_date.tzinfo is None:
            r.due_date=r.due_date.replace(tzinfo=tz)
        if r.due_date<datetime.now(tz) and r.status==Status.pending:
            Question.delete_expired(r)
        
    response={"user_name":user.username,"assessments":[{"id":r.id,"subject":r.subject,"topic":r.topic,"status":r.status.value,"created_at":r.time,"due_date":r.due_date, "score":r.score}  for r in res if r.due_date.replace(tzinfo=tz)>datetime.now(tz) ]}
    return jsonify(response)

# deleting the account
@routes.route("/delete",methods=["POST"])
def delete():
    details=request.json
    user_id=details["user_id"]
    user=User.query.filter_by(id=user_id).first()
    try:
        if  user:
            db.session.delete(user)
            db.session.commit()
    except Exception as e:
        return jsonify({"message":"Something went wrong!!! Try again"}),400
    return jsonify(),200

    

#submitted answers
@routes.route("/submitting",methods=["POST"])
#@cross_origin(origin='http://localhost:8080')
def submitting():
    data=request.json
    user_id=data["user_id"]
    score_id=data["score_id"]
    questions=data["answers"]
    print(type(questions))
    for qid_str , choice in questions.items():
        qid=int(qid_str)
        rec=Question.query.get(qid)
        rec.user_choice=choice
    sc_cal=Question.query.filter_by(score_id=score_id).all()
    count=0
    for i in sc_cal:
        if(i.user_choice==i.is_correct):
            count+=1
    score=Score.query.filter_by(user_id=user_id,id=score_id).first()
    if(score):
        score.score=count
    db.session.commit()
    return jsonify({"message":"choices added successfully","score":count})

# code generator
@routes.route("/code_generator",methods=["POST"])
def code_generator():
    data=request.json
    query=data["query"]
    language=data["language"]
    api_key=Config.API_KEY1
    try:
        genai.configure(api_key=api_key)
        # api model
        model = genai.GenerativeModel(
        model_name="gemini-2.0-flash",
        generation_config={
        "temperature": 0.2,
        "top_p": 1,
        "top_k": 1,
        "max_output_tokens": 2048,
        })
        code_prompt = f"""
                    Generate the code in {language} for the following task: {query}.
                     Include in-line comments in the code (explain key steps briefly).
                     Do NOT use Markdown formatting.
                     Return a proper explanation of the code in plain text (no markdown).
                     Respond strictly in JSON format like this:
                     {{
                     "code":"print(\\"hello\\"),
                     "explanation":""print is a built-in Python function. It's used to send output to the standard output device, usually the console or terminal. 
                     \\"hello\\" is a string literal, i.e., a sequence of characters surrounded by quotes (\\\" \\\" or '\\' '). 
                     Here, it's the argument to the print() function. Python passes \\"hello\\" into the print() function to be displayed.""
                     }} 
                     Only return a valid JSON object. Do not include anything outside the JSON."""
        code_response = model.generate_content(code_prompt)
        #cleaning the request to suitable json format
        code = re.sub(r"```json\n|\n```", "",code_response.text).strip()
        #converting to dict
        code_json=json.loads(code)
        
    except Exception as e:
        return jsonify({"message":"{e},something went wrong"}),400
    return jsonify({"code":code_json["code"],"explanation":code_json["explanation"]}),200
                    
# recent activity
@routes.route("/recent_activity", methods=["POST"])
def recent_activity():
    try:
        data = request.json
        user_id = data["user_id"]
        threshold = datetime.now(tz) - timedelta(days=7)
        #sort by time descending
        user = (
            Score.query
            .filter(Score.user_id == user_id, Score.time >= threshold)
            .order_by(Score.time.desc())
            .all()
        )
        response = []
        for u in user:
            # due days for pending assessments
            time = (u.due_date.date() - datetime.today().date()).days
            delta = (datetime.now(tz) - u.time) if u.time.tzinfo else (datetime.now(tz) - u.time.replace(tzinfo=tz))
            if delta.days < 1:
                settime = f"{delta.total_seconds()/3600:.1f} hours ago"
            else:
                settime = f"{u.time.strftime('%Y-%m-%d')} at {u.time.strftime('%H:%M')}"
            if u.status == Status.completed:
                response.append({
                    "id": u.id,
                    "title": f"{u.subject} Assessment completed ",
                    "time": settime,
                    "status": u.status.value,
                    "description": f" scored {round(((u.score/30)*100),2)}% on {u.topic}"
                })
            elif u.status == Status.pending:
                response.append({
                    "id": u.id,
                    "title": f"{u.subject} Assessment assigned ",
                    "time": settime,
                    "status": u.status.value,
                    "description": f"{u.topic} due in {time} days"
                })
            else:
                response.append({
                    "id": u.id,
                    "title": f"{u.subject} Assessment expired",
                    "time": settime,
                    "status": u.status.value,
                    "description": f"Assessment expired at {u.due_date.date().strftime('%Y-%m-%d')}"
                })
        return jsonify(response), 200
    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500
    
# Total assessment
@routes.route("/total_assessment", methods=["POST"])
def total_assessment():
    try:
        data = request.json
        user_id = data["user_id"]
        this_month = datetime.now(tz).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        now = datetime.now(tz)

        # Current month stats
        current_stats = (
            db.session.query(
                func.count(Score.id),
                func.avg(Score.score)
            )
            .filter(
                Score.user_id == user_id,
                Score.time >= this_month,
                Score.time <= now
            )
            .first()
        )
        count = current_stats[0] or 0
        average_score = int(current_stats[1]) if current_stats[1] is not None else 0

        # Previous month stats
        last_day_prev_month = this_month - timedelta(days=1)
        first_day_prev_month = last_day_prev_month.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        prev_stats = (
            db.session.query(
                func.count(Score.id),
                func.avg(Score.score)
            )
            .filter(
                Score.user_id == user_id,
                Score.time >= first_day_prev_month,
                Score.time <= last_day_prev_month
            )
            .first()
        )
        previouscount = prev_stats[0] or 0
        average_prev_month = int(prev_stats[1]) if prev_stats[1] is not None else 0

        # Calculate percentages
        days_in_current = (now - this_month).days + 1
        days_in_prev = (last_day_prev_month - first_day_prev_month).days + 1

        previous_percent = (previouscount / (days_in_prev * 3) * 100) if previouscount else 0
        current_percent = (count / (days_in_current * 3) * 100) if count else 0
        total_percent = current_percent - previous_percent
        approximate_avg_score = average_score - average_prev_month
        
        return jsonify({
            "total_assessment": count,
            "percentage": f"{total_percent:.2f}%",
            "average": average_score,
            "approxi_average": approximate_avg_score
        }), 200
    

    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500

# analysis of the assessment
@routes.route("/analysis",methods=["POST"])
def analysis():
    try:
        data=request.json
        user_id=data["user_id"]
       
        user=Score.query.filter_by(user_id=user_id).all()
        if not user:
            return jsonify({"message":"No assessments found"}),404
    
        year=data["selectedYear"]
        year_start=datetime(year,1,1,0,0,0,0)
        year_end=datetime(year,12,31,23,59,59,999999)
        

        # nothing available in the specific year
        if not any(score.time >= year_start and score.time <= year_end for score in user):
            return jsonify({"message": "No assessments found for the specified year"}), 404    
        topic_scores=[]
        # month wise subject average
        
        month=data["selectedMonth"]      
        month_start=datetime(year,month,1)
        last_day=monthrange(year,month)[1]   
        month_end= datetime(year,month,last_day,23,59,59,999999) 
        # topic wise scores
        topic_stat=(
            db.session.query(
                Score.topic,
                Score.time,
                Score.score,
                Score.difficulty
            ).filter(
                Score.user_id == user_id,
                Score.time >= month_start,
                Score.time <= month_end
            ).order_by(
                Score.topic, 
            )
        )
        topic_scores = [{'topic': topic, 'date':time.strftime("%Y-%m-%d"),'score': score if score is not None else 0,"difficulty":difficulty.value} 
                        for topic,time,score,difficulty in topic_stat]      
        return jsonify(topic_scores), 200
       
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500        
    
    
#subject wise analysis
@routes.route("/sub_analysis",methods=["POST"])
def sub_analysis():
    try:
        data=request.json
        user_id=data["user_id"]
       
        user=Score.query.filter_by(user_id=user_id).all()
        if not user:
            return jsonify({"message":"No assessments found"}),404
    
        year=data["selectedYear"]
        year_start=datetime(year,1,1,0,0,0,0)
        year_end=datetime(year,12,31,23,59,59,999999)
        

        # nothing available in the specific year
        if not any(score.time >= year_start and score.time <= year_end for score in user):
            return jsonify({"message": "No assessments found for the specified year"}), 404    
        subject_scores = []
        # month wise subject average
        
        month=data["selectedMonth"]      
        month_start=datetime(year,month,1)
        last_day=monthrange(year,month)[1]   
        month_end= datetime(year,month,last_day,23,59,59,999999) 
        subject_stats=(
            db.session.query(
                Score.subject,
                extract('month',Score.time).label('month'), 
                func.avg(Score.score).label('average_score')
            ).filter(
                Score.user_id == user_id,
                Score.time >= month_start,
                Score.time <= month_end
            ).group_by(
                Score.subject, extract('month',Score.time)
            
            ).order_by(
                Score.subject          
            )
        )
        
        subject_scores = [
                {"month": month_abbr[month], "subject": subject, "average_score": average_score if average_score is not None else 0} 
                  for subject, month, average_score in subject_stats]
        return jsonify(subject_scores), 200
           
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500
    
# for performance analysis
@routes.route("/performance_analysis",methods=["POST"])
def performance_analysis():
    try:
        data=request.json
        user_id=data["user_id"]
       
        user=Score.query.filter_by(user_id=user_id).all()
        if not user:
            return jsonify({"message":"No assessments found"}),404
    
        year=data["selectedYear"]
        year_start=datetime(year,1,1,0,0,0,0)
        year_end=datetime(year,12,31,23,59,59,999999)
        

        # nothing available in the specific year
        if not any(score.time >= year_start and score.time <= year_end for score in user):
            return jsonify({"message": "No assessments found for the specified year"}), 404    
        average_scores=[]
        # month wise average scores
        average_stats=(
        db.session.query(
            extract('month',Score.time).label('month'),
            func.avg(Score.score).label('average_score'),
            func.count(Score.id).label('total_assessments')
        ) .filter(
                Score.user_id == user_id,
                Score.time >= year_start,
                Score.time <= year_end
        ).group_by('month').order_by('month').all()
        )  
         
        average_scores= [
            {"month": month_abbr[month], "average_score": average_score if average_score is not None else 0, "total_assessments": total_assessments}  
            for month, average_score, total_assessments in average_stats] 
        return jsonify(average_scores), 200
       
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500
    
# Preview of the assessment
@routes.route("/preview",methods=["POST"])
def preview():
    data=request.json
    user_id=data.get("user_id")
    id=data.get("score_id")

    #storing result
    score=Score.query.filter_by(user_id=user_id,id=id).first()
    if score is None:
        return jsonify({"error": "Assessment not found"}), 404
    result={"id":id,"subject":score.subject,"title":score.topic,"questions":[]}
    question=Question.query.filter_by(score_id=id)
    for q in question:
        options = []
        for c in q.choices:
            options.append({"id":c["id"],"option":c["text"]})
        question_data = {
            "id":q.id,
            "text": q.quest_text,
            "options": options,
            "correct_option":q.is_correct,
            "user_choice":q.user_choice
        }
        result["questions"].append(question_data)
    
    return jsonify(result)

 # Update the status of the score
@routes.route("/update_exam_status", methods=["POST"])
def update_exam_status():
    data = request.json
    user_id = data.get("user_id")
    score_id = data.get("score_id")
    status = data.get("status")

    if not user_id or not score_id or not status:
        return jsonify({"error": "Missing required fields"}), 400

    score = Score.query.filter_by(user_id=user_id, id=score_id).first()
    if not score:
        return jsonify({"error": "Score not found"}), 404

    score.status = Status[status]

# profile 
@routes.route("/profile",methods=["POST"])
def profile():
    data=request.json
    user_id=data.get("user_id")
    user=User.query.filter_by(id=user_id).first()
    return jsonify({"user":{"email":user.email,"name":user.username}}),200
    


    



    



    

    